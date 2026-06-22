import zipfile
import xml.etree.ElementTree as ET
import os

def extract_docx_text(docx_path, output_path):
    print(f"Opening docx file: {docx_path}")
    try:
        # First try importing python-docx in case it's installed
        import docx
        doc = docx.Document(docx_path)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        # also tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                fullText.append(" | ".join(row_text))
        text = "\n".join(fullText)
        with open(output_path, "w", encoding="utf-8") as f:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
        print("Successfully extracted text using python-docx.")
        return
    except ImportError:
        print("python-docx not found. Falling back to robust pure XML parser...")
    
    # Robust fallback using zipfile and xml parsing (no dependencies)
    try:
        with zipfile.ZipFile(docx_path) as z:
            # Parse main document XML
            doc_xml = z.read("word/document.xml")
            root = ET.fromstring(doc_xml)
            
            # XML namespace map
            ns = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            text_parts = []
            # Find all paragraph elements
            for paragraph in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                para_text = []
                # Find all text elements inside paragraph
                for text_elem in paragraph.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if text_elem.text:
                        para_text.append(text_elem.text)
                text_parts.append("".join(para_text))
                
            full_text = "\n".join(text_parts)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(full_text)
            print(f"Successfully extracted text to: {output_path}")
    except Exception as e:
        print(f"Error during fallback extraction: {e}")

if __name__ == "__main__":
    docx_path = r"E:\Chao的资料\AI 方案\new idea\网站架构V1.2.docx"
    output_path = r"e:\CraftonAI\extracted_docx_v12_text.txt"
    extract_docx_text(docx_path, output_path)
